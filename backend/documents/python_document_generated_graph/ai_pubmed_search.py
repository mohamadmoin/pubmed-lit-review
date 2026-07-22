#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
AI-Powered PubMed Search and Content Generation

This module provides functionality for:
1. AI-guided document structure planning
2. PubMed search with NLP-based filtering
3. Full text retrieval and analysis
4. AI-powered content generation with citations
"""

import ast
import os
import re
import logging
import time
import xml.etree.ElementTree as ET
import json
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
import numpy as np
from pydantic import BaseModel
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from pymed import PubMed
import requests
from tqdm import tqdm
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .engine.entrez_client import EntrezClient
from .pubmed_structured_retrieval import (
    Article,
    FullTextArticle,
    fetch_full_text,
    lookup_pmc_ids_for_pmids,
    parse_pubmed_xml,
    parse_pmc_xml,
)
from .text_budget import (
    format_paper_for_selection,
    llm_worker_count,
    prepare_paper_content_for_summary,
    SELECTION_MAX_CANDIDATES,
    SELECTION_OUTPUT_MAX_TOKENS,
    SECTION_CONTENT_OUTPUT_MAX_TOKENS,
    SECTION_MAX_PAPERS_PER_CALL,
    SUMMARY_OUTPUT_MAX_TOKENS,
    truncate_summary_for_section,
)


def _get_llm_model():
    try:
        from litreview.llm_client import get_llm_model
        return get_llm_model()
    except Exception:
        return "gpt-4o-mini"


def _chat_completions_create(client, **kwargs):
    try:
        from litreview.llm_client import chat_completions_create
        kwargs.pop('model', None)
        return chat_completions_create(**kwargs)
    except Exception:
        if 'model' not in kwargs:
            kwargs['model'] = _get_llm_model()
        return client.chat.completions.create(**kwargs)


def _get_message_content(response):
    try:
        from litreview.llm_client import get_message_content
        return get_message_content(response.choices[0].message)
    except Exception:
        return (response.choices[0].message.content or '').strip()


def _parse_json_response(text: str):
    try:
        from litreview.llm_client import parse_json_response
        return parse_json_response(text)
    except Exception:
        content = text.strip()
        if content.startswith('```'):
            content = re.sub(r'^```(?:json)?\s*', '', content)
            content = re.sub(r'\s*```$', '', content)
        return json.loads(content)


def _normalize_document_structure(data: dict, subject: str) -> dict:
    """Normalize LLM output to the schema expected by DocumentStructure."""
    sections_in = data.get('sections') or []
    normalized_sections = []
    for section in sections_in:
        if not isinstance(section, dict):
            continue
        title = section.get('title') or section.get('heading') or section.get('name') or 'Section'
        description = (
            section.get('description')
            or section.get('content_summary')
            or section.get('content')
            or ''
        )
        search_terms = section.get('search_terms') or section.get('keywords') or []
        if isinstance(search_terms, str):
            search_terms = [search_terms]
        if not search_terms:
            search_terms = [title, subject]
        filter_keywords = section.get('filter_keywords') or search_terms[:3]
        if isinstance(filter_keywords, str):
            filter_keywords = [filter_keywords]
        normalized_sections.append({
            'title': title,
            'description': description,
            'search_terms': search_terms,
            'filter_keywords': filter_keywords,
        })
    if not normalized_sections:
        normalized_sections = [{
            'title': 'Overview',
            'description': data.get('description') or subject,
            'search_terms': [subject, f'{subject} review'],
            'filter_keywords': [subject.split()[0] if subject else 'research'],
        }]
    data['sections'] = normalized_sections
    data.setdefault('title', subject)
    data.setdefault('description', data.get('description') or subject)
    data.setdefault('citation_style', 'vancouver')
    return data

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(name)s - %(message)s',
    handlers=[
        logging.FileHandler("ai_pubmed_search.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("AIPubMedSearch")

# =============================================================================
# Data Classes
# =============================================================================

class DocumentStructure(BaseModel):
    """Represents the structure of a document to be generated."""
    title: str
    description: str
    sections: List[Dict[str, Any]]  # List of section details with search terms
    citation_style: str = "vancouver"  # Default citation style

class SearchResult(BaseModel):
    """Represents a search result with metadata and content."""
    
    title: str
    authors: str
    journal: str
    publication_date: Any  # Changed from str to Any to accept datetime.date
    abstract: str
    full_text: Optional[str] = None
    pmc_id: Optional[str] = None
    full_text_available: bool = False
    pmid: Optional[str]
    relevance_score: float = 0.0
    section_relevance: Dict[str, float] = field(default_factory=dict)
    search_term: Optional[str] = None
    section_title: Optional[str] = None

class ContentSection(BaseModel):
    """Represents a section of generated content with citations."""
    title: str
    content: str
    citations: List[Tuple[str, str, str, str, str]]  # List of (pmid, title, authors, journal, date) tuples
    ai_generated: bool = False  # Whether this section was AI-generated

class ProcessLogger:
    """Handles logging of process steps to Excel file."""
    def __init__(self, output_dir: str = "logs"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.excel_path = os.path.join(output_dir, f"process_log_{self.timestamp}.xlsx")
        self.writer = pd.ExcelWriter(self.excel_path, engine='openpyxl')
        
        # Initialize DataFrames for different sheets
        self.steps_df = pd.DataFrame(columns=['Timestamp', 'Step', 'Status', 'Details'])
        self.papers_df = pd.DataFrame(columns=['PMID', 'Title', 'Authors', 'Journal', 'Date', 'Abstract', 'Full Text', 'Relevance Score'])
        self.sections_df = pd.DataFrame(columns=['Title', 'Content', 'Citations', 'AI Generated', 'Search Terms', 'Filter Keywords'])
        self.errors_df = pd.DataFrame(columns=['Timestamp', 'Step', 'Error Message', 'Details'])
        self.section_relevance_df = pd.DataFrame(columns=['PMID', 'Section', 'Relevance Score'])
        
        # Save initial DataFrames
        self.steps_df.to_excel(self.writer, sheet_name='Process Steps', index=False)
        self.papers_df.to_excel(self.writer, sheet_name='Papers', index=False)
        self.sections_df.to_excel(self.writer, sheet_name='Sections', index=False)
        self.errors_df.to_excel(self.writer, sheet_name='Errors', index=False)
        self.section_relevance_df.to_excel(self.writer, sheet_name='Section Relevance', index=False)
    
    def log_step(self, step: str, status: str, details: str = ""):
        """Log a process step."""
        self.steps_df = pd.concat([self.steps_df, pd.DataFrame([{
            'Timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'Step': step,
            'Status': status,
            'Details': details
        }])], ignore_index=True)
        self.steps_df.to_excel(self.writer, sheet_name='Process Steps', index=False)
    
    def log_paper(self, paper: SearchResult):
        """Log a paper's details."""
        # Log main paper details
        self.papers_df = pd.concat([self.papers_df, pd.DataFrame([{
            'PMID': paper.pmid,
            'Title': paper.title,
            'Authors': paper.authors,
            'Journal': paper.journal,
            'Date': paper.publication_date,
            'Abstract': paper.abstract,
            'Full Text': paper.full_text[:1000] + "..." if paper.full_text else None,
            'Relevance Score': paper.relevance_score
        }])], ignore_index=True)
        self.papers_df.to_excel(self.writer, sheet_name='Papers', index=False)
        
        # Log section-specific relevance scores
        for section, score in paper.section_relevance.items():
            self.section_relevance_df = pd.concat([self.section_relevance_df, pd.DataFrame([{
                'PMID': paper.pmid,
                'Section': section,
                'Relevance Score': score
            }])], ignore_index=True)
        self.section_relevance_df.to_excel(self.writer, sheet_name='Section Relevance', index=False)
    
    def log_section(self, section: ContentSection, search_terms: List[str] = None, filter_keywords: List[str] = None):
        """Log a section's details."""
        self.sections_df = pd.concat([self.sections_df, pd.DataFrame([{
            'Title': section.title,
            'Content': section.content,
            'Citations': json.dumps(section.citations),
            'AI Generated': section.ai_generated,
            'Search Terms': json.dumps(search_terms) if search_terms else None,
            'Filter Keywords': json.dumps(filter_keywords) if filter_keywords else None
        }])], ignore_index=True)
        self.sections_df.to_excel(self.writer, sheet_name='Sections', index=False)
    
    def log_error(self, step: str, error: Exception, details: str = ""):
        """Log an error."""
        self.errors_df = pd.concat([self.errors_df, pd.DataFrame([{
            'Timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'Step': step,
            'Error Message': str(error),
            'Details': details
        }])], ignore_index=True)
        self.errors_df.to_excel(self.writer, sheet_name='Errors', index=False)
    
    def close(self):
        """Close the Excel writer."""
        self.writer.close()

# =============================================================================
# Main Functions
# =============================================================================

def get_document_structure(
    subject: str,
    description: str,
    num_words: int,
    client: Any,
    use_enhanced_filtering: bool = True
) -> DocumentStructure:
    """
    Get AI-generated document structure and search terms.
    
    Args:
        subject: Main subject of the document
        description: Detailed description of what needs to be written
        num_words: Target number of words
        client: OpenAI client instance
        use_enhanced_filtering: Whether to use enhanced NLP filtering
        
    Returns:
        DocumentStructure object with planned sections and search terms
    """
    prompt = f"""
    Create a detailed document structure for a {num_words}-word document on {subject}.
    
    Description: {description}
    
    Return ONLY a JSON object (no markdown) with exactly these fields:
    - title: Main document title
    - description: Overall document description
    - sections: List of 2-4 section objects, each with:
        - title: Section title
        - description: Section description
        - search_terms: List of 3-5 PubMed search terms for this section
        - filter_keywords: List of 2-3 keywords for filtering results
    - citation_style: Preferred citation style (e.g. vancouver, apa)
    """
    
    try:
        response = _chat_completions_create(
            client,
            messages=[
                {"role": "system", "content": "You are a research document planning expert. Return valid JSON only."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.3,
            max_tokens=2048,
        )
        
        response_content = _get_message_content(response)
        if not response_content:
            raise ValueError("LLM returned empty content for document structure")
        
        structure_data = _parse_json_response(response_content)
        structure_data = _normalize_document_structure(structure_data, subject)
        
        return DocumentStructure(**structure_data)
    
    except Exception as e:
        logger.error(f"Error getting document structure: {e}")
        if 'response_content' in locals():
            logger.error(f"Raw LLM response (first 500 chars): {response_content[:500]!r}")
        raise


def search_pubmed_with_filtering(
    document_structure: DocumentStructure,
    max_results: int = 200,
    use_enhanced_filtering: bool = True
) -> List[SearchResult]:
    """
    Search PubMed and filter results using NLP for each section.
    
    Args:
        document_structure: DocumentStructure object with section-specific search terms
        max_results: Maximum results per search term
        use_enhanced_filtering: Whether to use enhanced NLP filtering
        
    Returns:
        List of filtered SearchResult objects
    """
    from concurrent.futures import ThreadPoolExecutor, as_completed
    from collections import defaultdict
    
    # Initialize TF-IDF vectorizer
    vectorizer = TfidfVectorizer()
    
    # Dictionary to store papers by section and keyword
    section_papers = defaultdict(lambda: defaultdict(list))
    
    def fetch_pubmed_results(search_term, section_title):
        pubmed = PubMed(tool="MyTool", email="my@email.address")
        max_retries = 3
        for attempt in range(max_retries):
            try:
                print(f"Searching for term: {search_term} in section: {section_title}")
                search_results = pubmed.query(search_term, max_results=20)
                
                # Convert iterator to list immediately
                search_results_list = list(search_results)
                print(f"Number of papers found: {len(search_results_list)}")
                
                # Process and return the list
                processed_results = []
                for article in search_results_list:
                    if hasattr(article, 'journal'):
                        # Convert date to string if it's a datetime.date object
                        pub_date = str(article.publication_date) if hasattr(article, 'publication_date') else "Unknown"
                        
                        # Handle missing abstract
                        abstract = article.abstract if hasattr(article, 'abstract') and article.abstract else ""
                        try:
                            aUthors = ', '.join([str(autho['lastname']) + " " + str(autho['initials']) for autho in article.authors[:6]])
                        except Exception as e:
                            aUthors = ""
                            
                        # Handle multiple PMIDs by taking only the first one
                        pmid = article.pubmed_id.split('\n')[0] if article.pubmed_id else None
                        result = SearchResult(
                            pmid=pmid,
                            title=article.title,
                            authors= aUthors,
                            journal=article.journal,
                            publication_date=article.publication_date,
                            abstract=abstract,
                            section_relevance={section_title: 0.0},  # Initialize section relevance
                            search_term=search_term,  # Store the search term that found this paper
                            section_title=section_title  # Store the section this paper belongs to
                        )
                        processed_results.append(result)
                
                return processed_results
                
            except requests.exceptions.ConnectionError as e:
                print(f"ConnectionError on attempt {attempt + 1}: {e}")
                time.sleep(5)  # Wait 5 seconds before retrying
            except Exception as e:
                print(f"An error occurred: {e}")
                break  # Exit if it's a different kind of error
        return []
    
    def process_search_term(args):
        """Helper function for parallel processing"""
        section_title, search_term = args
        results = fetch_pubmed_results(search_term, section_title)
        return section_title, search_term, results
    
    # Create a list of all search terms with their sections
    search_tasks = []
    for section in document_structure.sections:
        section_title = section['title']
        for term in section['search_terms']:
            search_tasks.append((section_title, term))
    
    # Process search terms in parallel using ThreadPoolExecutor
    with ThreadPoolExecutor(max_workers=3) as executor:
        future_to_task = {executor.submit(process_search_term, task): task for task in search_tasks}
        
        for future in as_completed(future_to_task):
            section_title, search_term, results = future.result()
            section_papers[section_title][search_term] = results
    
    # Process and filter results for each section and keyword combination
    final_results = []
    final_results_pre_filtered = []
    
    for section in document_structure.sections:
        section_title = section['title']
        
        # Process each keyword's papers separately
        for search_term, papers in section_papers[section_title].items():
            if not papers:
                continue
            
            # Create DataFrame for this keyword's papers
            df = pd.DataFrame([r.dict() for r in papers])
            
            if use_enhanced_filtering:
                # Create TF-IDF matrix for this keyword's papers
                df['combined_text'] = df['title'] + ' ' + df['abstract'].fillna('')
                tfidf_matrix = vectorizer.fit_transform(df['combined_text'])
                
                # Calculate similarity score with the search term that found these papers
                keyword_vec = vectorizer.transform([search_term])
                similarities = cosine_similarity(keyword_vec, tfidf_matrix).flatten()
                
                # Update similarity scores
                df['similarity_score'] = similarities
                
                # Sort by similarity score and apply threshold
                threshold = 0.05
                df_pre_filtered = df
                df = df[df['similarity_score'] > threshold]
                df = df.sort_values('similarity_score', ascending=False)
                
                # Keep top 10 papers
                df = df.head(10)
            
            # Convert filtered results back to SearchResult objects
            for _, row in df.iterrows():
                result_dict = row.to_dict()
                result_dict['section_relevance'] = {section_title: result_dict.get('similarity_score', 0.0)}
                result = SearchResult(**result_dict)
                final_results.append(result)
                
            for _, row in df_pre_filtered.iterrows():
                result_dict_pre_filtered = row.to_dict()
                result_dict_pre_filtered['section_relevance'] = {section_title: result_dict_pre_filtered.get('similarity_score', 0.0)}
                result_pre_filtered = SearchResult(**result_dict_pre_filtered)
                final_results_pre_filtered.append(result_pre_filtered)
    
    return final_results, final_results_pre_filtered


def _extract_paper_indices(response_text: str, max_index: int) -> List[int]:
    """Extract paper index list from LLM output (JSON, reasoning text, or prose)."""
    if not response_text or max_index <= 0:
        return []

    text = response_text.strip()

    # Direct JSON array
    start = text.find('[')
    end = text.rfind(']')
    if start != -1 and end > start:
        list_str = text[start:end + 1]
        try:
            parsed = ast.literal_eval(list_str)
            if isinstance(parsed, list):
                return [i for i in parsed if isinstance(i, int) and 0 <= i < max_index]
        except (SyntaxError, ValueError):
            pass

    # Comma-separated integers after keywords like "indices", "selected", "return"
    keyword_match = re.search(
        r'(?:indices|selected|papers?|return)\s*[:\-]?\s*\[?([\d,\s]+)\]?',
        text,
        re.IGNORECASE,
    )
    if keyword_match:
        nums = [int(n) for n in re.findall(r'\d+', keyword_match.group(1))]
        valid = [i for i in nums if 0 <= i < max_index]
        if valid:
            return valid[:5]

    # Numbered paper lines marked as high relevance / selected
    selected = []
    for match in re.finditer(
        r'(?:^|\n)\s*(?:\*|\-)?\s*(?:Paper\s*)?(\d+)[.:)]\s*(?:\*|Selected|High relevance|Relevance:\s*High)',
        text,
        re.IGNORECASE,
    ):
        idx = int(match.group(1))
        if 0 <= idx < max_index:
            selected.append(idx)
    if selected:
        return list(dict.fromkeys(selected))[:5]

    # Any standalone small integers that look like indices (last resort)
    candidates = [int(n) for n in re.findall(r'\b(\d+)\b', text) if int(n) < max_index]
    if candidates:
        return list(dict.fromkeys(candidates))[:5]

    return []


def select_relevant_papers(
    results: List[SearchResult],
    document_structure: DocumentStructure,
    client: Any
) -> Dict[str, List[SearchResult]]:
    """
    Use AI to select the most relevant papers for each section.
    
    Args:
        results: List of SearchResult objects
        document_structure: DocumentStructure object
        client: OpenAI client instance
        
    Returns:
        Dictionary mapping section titles to lists of selected SearchResult objects
    """
    # Group papers by section
    section_papers = {}
    for paper in results:
        section_title = paper.section_title
        if section_title not in section_papers:
            section_papers[section_title] = []
        section_papers[section_title].append(paper)
    
    # Dictionary to store selected papers by section
    selected_papers_by_section = {}
    
    # Process each section separately
    for section in document_structure.sections:
        section_title = section['title']
        section_description = section['description']
        
        # Get papers for this section
        section_results = section_papers.get(section_title, [])
        if not section_results:
            logger.warning(f"No papers found for section: {section_title}")
            selected_papers_by_section[section_title] = []
            continue
        
        # Remove duplicates based on PMID
        seen_pmids = set()
        unique_papers = []
        for paper in section_results:
            if paper.pmid not in seen_pmids:
                seen_pmids.add(paper.pmid)
                unique_papers.append(paper)
        
        # Sort by relevance and cap candidates to stay within context window
        unique_papers.sort(key=lambda p: getattr(p, 'relevance_score', 0) or 0, reverse=True)
        unique_papers = unique_papers[:SELECTION_MAX_CANDIDATES]

        papers_text = "\n".join([
            format_paper_for_selection(i, r)
            for i, r in enumerate(unique_papers)
        ])
        
        prompt = f"""
        Select the most relevant papers for the following section of the document.
        
        Section Title: {section_title}
        Section Description: {section_description}
        
        Available papers:
        {papers_text}
        
        Select the most relevant papers (maximum 5) and provide their indices.
        Return ONLY a JSON array of integers, e.g. [0, 1, 2]. No explanation.
        """
        
        try:
            response = _chat_completions_create(
                client,
                messages=[
                    {"role": "system", "content": "You select papers for literature reviews. Output ONLY a JSON array of integer indices, like [0,2,4]."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=SELECTION_OUTPUT_MAX_TOKENS,
            )
            
            response_text = _get_message_content(response)
            selected_indices = _extract_paper_indices(response_text, len(unique_papers))

            if not selected_indices:
                logger.warning(f"No list found in response: {response_text[:500]}")
                
            # Validate the response structure
            valid_indices = [i for i in selected_indices if isinstance(i, int) and 0 <= i < len(unique_papers)]
            
            if not valid_indices:
                logger.error(f"No valid paper indices found for section {section_title}; using top relevance scores")
                valid_indices = list(range(min(5, len(unique_papers))))
            
            # Get selected papers
            selected_papers = [unique_papers[i] for i in valid_indices]
            selected_papers_by_section[section_title] = selected_papers
            
        except Exception as e:
            logger.error(f"Error selecting papers for section {section_title}: {e}")
            # Fallback to first 5 papers
            selected_papers_by_section[section_title] = unique_papers[:5]
    
    return selected_papers_by_section

def retrieve_full_text(
    papers_by_section: Dict[str, List[SearchResult]],
    entrez_client: Optional[Any] = None
) -> Dict[str, List[SearchResult]]:
    """
    Retrieve full text for papers using Entrez API.

    Resolves PMID -> PMC ID before fetching from the PMC database.
    """
    if not entrez_client:
        from django.conf import settings as django_settings
        entrez_client = getattr(django_settings, 'ENTREZ_CLIENT', None)
        if not entrez_client:
            from .engine.entrez_client import EntrezClient
            from django.conf import settings as django_settings
            entrez_client = EntrezClient(
                tool=getattr(django_settings, 'PUBMED_TOOL', 'LitReview'),
                email=getattr(django_settings, 'PUBMED_EMAIL', None) or 'litreview@localhost',
                api_key=getattr(django_settings, 'PUBMED_API_KEY', None),
            )

    # Collect unique PMIDs across all sections
    pmid_to_paper: Dict[str, SearchResult] = {}
    for papers in papers_by_section.values():
        for paper in papers:
            if paper.pmid and paper.pmid not in pmid_to_paper:
                pmid_to_paper[paper.pmid] = paper

    pmc_map = lookup_pmc_ids_for_pmids(entrez_client, list(pmid_to_paper.keys()))
    logger.info(f"Resolved PMC IDs for {len(pmc_map)}/{len(pmid_to_paper)} papers")

    papers_with_text = {}

    for section_title, papers in papers_by_section.items():
        papers_with_text[section_title] = []

        for paper in tqdm(papers, desc=f"Retrieving full text for {section_title}"):
            try:
                if not paper.pmid:
                    papers_with_text[section_title].append(paper)
                    continue

                pmc_id = pmc_map.get(paper.pmid)
                paper.pmc_id = pmc_id
                paper.full_text_available = bool(pmc_id)

                if pmc_id:
                    full_text = fetch_full_text(entrez_client, pmc_id)
                    if full_text and full_text.raw_text:
                        paper.full_text = full_text.raw_text
                        paper.full_text_available = True
                        logger.info(f"Retrieved full text for PMID {paper.pmid} via {pmc_id}")
                    else:
                        logger.warning(
                            f"No full text content found for {pmc_id} (PMID {paper.pmid})"
                        )
                else:
                    logger.debug(f"No PMC open-access record for PMID {paper.pmid}")

                papers_with_text[section_title].append(paper)

            except Exception as e:
                logger.error(f"Failed to retrieve full text for PMID {paper.pmid}: {e}")
                papers_with_text[section_title].append(paper)

    return papers_with_text

def format_citation(title: str, authors: str, journal: str, date: str, citation_style: str = 'apa') -> str:
    """
    Format a citation according to the specified style.
    
    Args:
        title: Paper title
        authors: Author names
        journal: Journal name
        date: Publication date
        citation_style: Citation style (e.g., 'apa', 'mla', 'chicago')
        
    Returns:
        Formatted citation string
    """
    if citation_style == 'apa':
        # APA style: (Author, Year)
        year = date.split('-')[0] if date else 'n.d.'
        return f"({authors}, {year})"
    elif citation_style == 'mla':
        # MLA style: (Author page)
        return f"({authors})"
    elif citation_style == 'chicago':
        # Chicago style: (Author Year)
        year = date.split('-')[0] if date else 'n.d.'
        return f"({authors} {year})"
    else:
        # Default to APA style
        year = date.split('-')[0] if date else 'n.d.'
        return f"({authors}, {year})"

def format_reference(title: str, authors: str, journal: str, date: str, citation_style: str = 'apa') -> str:
    """
    Format a reference according to the specified style.
    
    Args:
        title: Paper title
        authors: Author names
        journal: Journal name
        date: Publication date
        citation_style: Citation style (e.g., 'apa', 'mla', 'chicago')
        
    Returns:
        Formatted reference string
    """
    if citation_style == 'apa':
        # APA style: Author, A. A., & Author, B. B. (Year). Title. Journal.
        year = date.split('-')[0] if date else 'n.d.'
        return f"{authors} ({year}). {title}. {journal}."
    elif citation_style == 'mla':
        # MLA style: Author, First. "Title." Journal, Date.
        return f"{authors}. \"{title}.\" {journal}, {date}."
    elif citation_style == 'chicago':
        # Chicago style: Author, First. "Title." Journal Date.
        return f"{authors}. \"{title}.\" {journal} {date}."
    else:
        # Default to APA style
        year = date.split('-')[0] if date else 'n.d.'
        return f"{authors} ({year}). {title}. {journal}."

def format_citations(sections: List[ContentSection], citation_style: str = 'apa') -> Tuple[List[ContentSection], str]:
    """
    Format citations in the content and create a references section.
    
    Args:
        sections: List of ContentSection objects
        citation_style: Citation style to use (e.g., 'apa', 'mla', 'chicago')
        
    Returns:
        Tuple of (updated sections, references text)
    """
    # Create a mapping of PMIDs to citation data
    pmid_to_citation = {}
    for section in sections:
        for pmid, title, authors, journal, date in section.citations:
            if pmid not in pmid_to_citation:
                pmid_to_citation[pmid] = (title, authors, journal, date)
    
    # Assign numbered references in order of first appearance in content
    pmid_to_number = {}
    next_number = 1
    citation_pattern = re.compile(r'\[\s*PMID:(\d+)\s*\]|PMID:(\d+)')
    for section in sections:
        for match in citation_pattern.finditer(section.content):
            pmid = match.group(1) or match.group(2)
            if pmid in pmid_to_citation and pmid not in pmid_to_number:
                pmid_to_number[pmid] = next_number
                next_number += 1
    for pmid in pmid_to_citation:
        if pmid not in pmid_to_number:
            pmid_to_number[pmid] = next_number
            next_number += 1

    # Replace PMID citations with numbered references [1], [2], ...
    for section in sections:
        content = section.content
        for match in citation_pattern.finditer(section.content):
            pmid = match.group(1) or match.group(2)
            if pmid in pmid_to_number:
                content = content.replace(
                    match.group(0),
                    f'[{pmid_to_number[pmid]}]',
                    1,
                )
        section.content = content
    
    # Generate references section with numbers
    references = []
    for pmid, number in sorted(pmid_to_number.items(), key=lambda item: item[1]):
        title, authors, journal, date = pmid_to_citation[pmid]
        reference = format_reference(title, authors, journal, date, citation_style)
        references.append(f"{number}. {reference}")
    
    references_text = '\n\n'.join(references)
    
    return sections, references_text

def summarize_paper(
    paper: SearchResult,
    client: Any
) -> str:
    """
    Generate a bullet-point summary of a paper using AI.
    
    Args:
        paper: SearchResult object containing paper details
        client: OpenAI client instance
        
    Returns:
        Bullet-point summary of the paper
    """
    # Prepare compact content for the local model context window
    paper_content = prepare_paper_content_for_summary(paper)
    
    prompt = f"""
    Create a focused bullet-point summary of this research paper for use in a literature review.
    
    Paper Details:
    Title: {paper.title}
    Authors: {paper.authors}
    Journal: {paper.journal}
    Publication Date: {paper.publication_date}
    
    Content:
    {paper_content}
    
    Include these points when available:
    1. Key findings and conclusions
    2. Methodology
    3. Important statistics or data
    4. Limitations or future directions
    
    Use concise bullet points (•). Maximum 8 bullets.
    """
    
    try:
        response = _chat_completions_create(
            client,
            messages=[
                {"role": "system", "content": "You summarize research papers concisely for literature reviews."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=SUMMARY_OUTPUT_MAX_TOKENS,
        )
        
        return _get_message_content(response)
    
    except Exception as e:
        logger.error(f"Error summarizing paper {paper.pmid}: {e}")
        abstract = (paper.abstract or '').strip()
        if abstract:
            return f"• Abstract-based summary (LLM unavailable):\n• {abstract[:800]}"
        return f"Summary unavailable for: {paper.title}"


def _build_section_papers_text(section_papers, paper_summaries) -> str:
    return "\n\n".join([
        f"[PMID:{p.pmid}]\nTitle: {p.title}\n"
        f"Authors: {p.authors}\n"
        f"Journal: {p.journal}\n"
        f"Date: {p.publication_date}\n"
        f"Search Term: {p.search_term}\n"
        f"Summary:\n{truncate_summary_for_section(paper_summaries.get(p.pmid))}\n"
        for p in section_papers
    ])


def _generate_section_content_chunk(
    section_title: str,
    section_desc: str,
    section_papers: List[SearchResult],
    paper_summaries: Dict[str, str],
    client: Any,
) -> str:
    """Generate section narrative for a subset of papers."""
    papers_text = _build_section_papers_text(section_papers, paper_summaries)
    prompt = f"""
    Write content for the following section based on the provided paper summaries.

    Section: {section_title}
    Description: {section_desc}

    Papers and their summaries:
    {papers_text}

    Write the content in a narrative style, citing papers using [PMID:X] format.
    If information is missing from the papers, use your knowledge but mark it as [AI_GENERATED].
    Synthesize the information from the papers into a coherent narrative.
    """
    response = _chat_completions_create(
        client,
        messages=[
            {"role": "system", "content": "You are a research content generation expert."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.5,
        max_tokens=SECTION_CONTENT_OUTPUT_MAX_TOKENS,
    )
    return _get_message_content(response)


def _generate_section_with_fallback(
    section_title: str,
    section_desc: str,
    section_papers: List[SearchResult],
    paper_summaries: Dict[str, str],
    client: Any,
) -> str:
    """Generate section content; split into chunks if the combined prompt is too large."""
    try:
        return _generate_section_content_chunk(
            section_title, section_desc, section_papers, paper_summaries, client
        )
    except Exception as exc:
        if len(section_papers) <= 1:
            raise
        logger.warning(
            f"Section '{section_title}' failed ({exc}); generating in smaller chunks"
        )

    chunks = []
    chunk_size = max(1, SECTION_MAX_PAPERS_PER_CALL)
    for i in range(0, len(section_papers), chunk_size):
        chunk_papers = section_papers[i:i + chunk_size]
        chunk_label = f"{section_title} (part {i // chunk_size + 1})"
        chunk_content = _generate_section_content_chunk(
            chunk_label, section_desc, chunk_papers, paper_summaries, client
        )
        chunks.append(chunk_content)

    if len(chunks) == 1:
        return chunks[0]

    merge_prompt = f"""
    Merge the following draft subsections into one cohesive section titled "{section_title}".
    Preserve [PMID:X] citations. Remove duplicate sentences.

    Drafts:
    {chr(10).join(f'--- Part {i+1} ---{chr(10)}{c}' for i, c in enumerate(chunks))}
    """
    response = _chat_completions_create(
        client,
        messages=[
            {"role": "system", "content": "You merge literature review subsections into one section."},
            {"role": "user", "content": merge_prompt},
        ],
        temperature=0.4,
        max_tokens=SECTION_CONTENT_OUTPUT_MAX_TOKENS,
    )
    return _get_message_content(response)


def generate_content(
    papers_by_section: Dict[str, List[SearchResult]],
    document_structure: DocumentStructure,
    client: Any,
    document_id: str,
    neo4j_client: Any
) -> List[ContentSection]:
    """
    Generate content for each section using AI with parallel processing.
    """
    from concurrent.futures import ThreadPoolExecutor, as_completed
    sections = []
    
    # First, generate summaries for all unique papers
    logger.info("Generating paper summaries...")
    paper_summaries = {}
    
    # First pass: collect all unique papers
    all_pmids = set()
    for section_papers in papers_by_section.values():
        for paper in section_papers:
            all_pmids.add(paper.pmid)
    
    # Helper function for parallel summary generation
    def generate_summary(pmid):
        # Find the first instance of this paper
        paper = next(
            paper for section_papers in papers_by_section.values()
            for paper in section_papers
            if paper.pmid == pmid
        )
        summary = summarize_paper(paper, client)
        return pmid, summary
    
    # Generate summaries in parallel
    workers = llm_worker_count()
    with ThreadPoolExecutor(max_workers=workers) as executor:
        future_to_pmid = {
            executor.submit(generate_summary, pmid): pmid 
            for pmid in all_pmids
        }
        
        for future in tqdm(
            as_completed(future_to_pmid), 
            total=len(all_pmids),
            desc="Generating paper summaries"
        ):
            pmid = future_to_pmid[future]
            try:
                pmid, summary = future.result()
                paper_summaries[pmid] = summary
            except Exception as e:
                logger.error(f"Error generating summary for PMID {pmid}: {e}")
                paper_summaries[pmid] = f"Error generating summary: {str(e)}"
    
    # Store paper summaries in Neo4j
    neo4j_client.store_paper_summaries(document_id, paper_summaries)
    
    # Helper function for parallel content generation
    def generate_section_content(section_info):
        section_title, section_desc, section_papers = section_info
        
        if not section_papers:
            logger.warning(f"No papers found for section: {section_title}")
            return None
        
        try:
            content = _generate_section_with_fallback(
                section_title,
                section_desc,
                section_papers,
                paper_summaries,
                client,
            )
            
            # Extract citations with full metadata
            citations = []
            for match in re.finditer(r'\[PMID:(\d+)\]', content):
                pmid = match.group(1)
                paper = next((p for p in section_papers if p.pmid == pmid), None)
                if paper:
                    pub_date = str(paper.publication_date) if hasattr(paper.publication_date, 'strftime') else str(paper.publication_date)
                    citations.append((
                        paper.pmid,
                        paper.title,
                        paper.authors,
                        paper.journal,
                        pub_date
                    ))
            
            return ContentSection(
                title=section_title,
                content=content,
                citations=citations,
                ai_generated='[AI_GENERATED]' in content
            )
            
        except Exception as e:
            logger.error(f"Error generating content for section {section_title}: {e}")
            return None
    
    # Prepare section generation tasks
    section_tasks = [
        (
            section['title'],
            section['description'],
            papers_by_section.get(section['title'], [])
        )
        for section in document_structure.sections
    ]
    
    # Generate content for sections in parallel
    workers = llm_worker_count()
    with ThreadPoolExecutor(max_workers=workers) as executor:
        future_to_section = {
            executor.submit(generate_section_content, task): task[0]  # section_title
            for task in section_tasks
        }
        
        for future in tqdm(
            as_completed(future_to_section),
            total=len(section_tasks),
            desc="Generating section content"
        ):
            section_title = future_to_section[future]
            try:
                section = future.result()
                if section:
                    sections.append(section)
            except Exception as e:
                logger.error(f"Error processing section {section_title}: {e}")
    
    # Sort sections to maintain original order
    sections.sort(key=lambda x: next(
        i for i, s in enumerate(document_structure.sections) 
        if s['title'] == x.title
    ))
    
    return sections

def create_word_document(
    sections: List[ContentSection],
    references: str,
    output_dir: str = "output",
    filename: str = None
) -> str:
    """
    Create a Word document from the generated content and references.
    
    Args:
        sections: List of ContentSection objects
        references: References section text
        output_dir: Directory to save the document
        filename: Optional filename (if None, generates one with timestamp)
        
    Returns:
        Path to the created Word document
    """
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate filename if not provided
    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"research_content_{timestamp}.docx"
    
    doc_path = os.path.join(output_dir, filename)
    
    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading(sections[0].title if sections else "Research Document", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add each section
    for section in sections:
        # Add section title
        doc.add_heading(section.title, level=1)
        
        # Add section content
        paragraph = doc.add_paragraph(section.content)
        paragraph.style.font.size = Pt(11)
        paragraph.style.font.name = 'Times New Roman'
        
        # Add spacing between sections
        doc.add_paragraph()
    
    # Add references section
    doc.add_heading("References", level=1)
    doc.add_paragraph(references)
    
    # Save the document
    doc.save(doc_path)
    
    return doc_path

def generate_research_content(
    subject: str,
    description: str,
    num_words: int,
    openai_client: Any,
    entrez_client: Optional[EntrezClient] = None,
    use_enhanced_filtering: bool = True,
    tool: str = "AIPubMedSearch",
    email: str = "your@email.com",
    output_filename: str = None
) -> Tuple[List[ContentSection], str, str]:
    """
    Main function to generate research content using AI and PubMed.
    
    Args:
        subject: Main subject of the research document
        description: Detailed description of the research document
        num_words: Target number of words in the document
        openai_client: OpenAI client instance for AI-powered content generation
        entrez_client: Optional EntrezClient instance for full text retrieval
        use_enhanced_filtering: Whether to use enhanced NLP filtering
        tool: Name of the tool for logging
        email: Email for logging
        output_filename: Optional filename for the Word document
        
    Returns:
        Tuple of (List of ContentSection objects, References section text, Path to Word document)
    """
    # Initialize process logger with a different name to avoid conflict
    process_logger = ProcessLogger()
    
    try:
        # Initialize EntrezClient if not provided
        if entrez_client is None:
            entrez_client = EntrezClient(tool=tool, email=email)
        
        # 1. Get document structure and search terms
        process_logger.log_step("Document Structure", "Started", f"Subject: {subject}")
        structure = get_document_structure(
            subject, description, num_words, openai_client, use_enhanced_filtering
        )
        process_logger.log_step("Document Structure", "Completed", json.dumps(structure.dict()))
        
        # 2. Search PubMed and filter results
        search_terms_summary = {
            section['title']: section['search_terms']
            for section in structure.sections
        }
        print("//////////////////////////////////////////////////////////////")
        print("//////////////////////////////////////////////////////////////")
        process_logger.log_step("PubMed Search", "Started", f"Section-specific search terms: {json.dumps(search_terms_summary)}")
        results, df_pre_filtered = search_pubmed_with_filtering(structure, use_enhanced_filtering=use_enhanced_filtering)
        process_logger.log_step("PubMed Search", "Completed", f"Found {len(results)} papers")
        print("//////////////////////////////////////////////////////////////")
        print("//////////////////////////////////////////////////////////////")
        for paper in results:
            process_logger.log_paper(paper)
        
        # 3. Select relevant papers
        process_logger.log_step("Paper Selection", "Started")
        selected_papers_by_section = select_relevant_papers(results, structure, openai_client)
        
        # Combine all selected papers into a single list for logging
        all_selected_papers = []
        for section_title, papers in selected_papers_by_section.items():
            process_logger.log_step("Paper Selection", "Section Completed", f"Selected {len(papers)} papers for section: {section_title}")
            for paper in papers:
                process_logger.log_paper(paper)
                all_selected_papers.append(paper)
        
        process_logger.log_step("Paper Selection", "Completed", f"Selected total of {len(all_selected_papers)} papers across all sections")
        
        # 4. Retrieve full text
        process_logger.log_step("Full Text Retrieval", "Started")
        papers_with_text = retrieve_full_text(selected_papers_by_section, entrez_client)
        process_logger.log_step("Full Text Retrieval", "Completed", f"Retrieved full text for {len(papers_with_text)} papers")
        
        # 5. Generate content
        process_logger.log_step("Content Generation", "Started")
        sections = generate_content(papers_with_text, structure, openai_client, "document_id", "neo4j_client")
        process_logger.log_step("Content Generation", "Completed", f"Generated {len(sections)} sections")
        for section in sections:
            process_logger.log_section(section)
        
        # 6. Format citations and create references
        process_logger.log_step("Citation Formatting", "Started")
        final_sections, references = format_citations(sections, structure.citation_style)
        process_logger.log_step("Citation Formatting", "Completed")
        
        # Create Word document
        process_logger.log_step("Document Creation", "Started")
        doc_path = create_word_document(final_sections, references, filename=output_filename)
        process_logger.log_step("Document Creation", "Completed", f"Document saved to: {doc_path}")
        
        process_logger.log_step("Process", "Completed", "All steps completed successfully")
        return final_sections, references, doc_path
        
    except Exception as e:
        process_logger.log_error("Main Process", e)
        raise
    finally:
        process_logger.close() 